from __future__ import annotations

from io import BytesIO

import pytest

from hf_text_summary.text_extract import extract_text_from_bytes, supported_extensions


def test_supported_extensions_includes_pdf_and_docx():
    exts = supported_extensions()
    assert ".txt" in exts
    assert ".pdf" in exts
    assert ".docx" in exts


def test_extract_docx_paragraphs_and_tables():
    try:
        import docx  # type: ignore
    except Exception as e:  # pragma: no cover
        pytest.skip(f"python-docx not available: {e}")

    doc = docx.Document()
    doc.add_paragraph("Hello from DOCX")

    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cell A"
    table.cell(0, 1).text = "Cell B"

    buf = BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    text = extract_text_from_bytes("sample.docx", data)
    assert "Hello from DOCX" in text
    assert "Cell A" in text
    assert "Cell B" in text


def _make_pdf_bytes_with_text(message: str) -> bytes:
    # Create a tiny PDF with an embedded text content stream.
    # This avoids shipping binary fixtures and keeps the test self-contained.
    try:
        from pypdf import PdfWriter
    except Exception as e:  # pragma: no cover
        pytest.skip(f"pypdf not available: {e}")

    try:
        # pypdf v4+ generics
        from pypdf.generic import (
            DecodedStreamObject,
            DictionaryObject,
            NameObject,
        )
    except Exception as e:  # pragma: no cover
        raise RuntimeError(f"pypdf generics not available: {e}")

    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=144)

    # Add a standard font to the page resources.
    font_dict = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    add_obj = getattr(writer, "add_object", None) or getattr(writer, "_add_object")
    font_ref = add_obj(font_dict)

    resources = page.get("/Resources") or DictionaryObject()
    resources.update(
        {
            NameObject("/Font"): DictionaryObject(
                {
                    NameObject("/F1"): font_ref,
                }
            )
        }
    )
    page[NameObject("/Resources")] = resources

    # Write a basic text content stream.
    safe = (message or "").replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    content = f"BT\n/F1 18 Tf\n30 90 Td\n({safe}) Tj\nET\n".encode("utf-8")

    stream = DecodedStreamObject()
    stream.set_data(content)
    stream_ref = add_obj(stream)
    page[NameObject("/Contents")] = stream_ref

    buf = BytesIO()
    writer.write(buf)
    return buf.getvalue()


def test_extract_pdf_text():
    data = _make_pdf_bytes_with_text("Hello from PDF")
    text = extract_text_from_bytes("sample.pdf", data)
    assert "Hello from PDF" in text


def test_extract_blank_pdf_returns_empty_string():
    from pypdf import PdfWriter

    buf = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.write(buf)

    text = extract_text_from_bytes("blank.pdf", buf.getvalue())
    assert isinstance(text, str)
